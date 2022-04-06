import sqlite3
import docx
import tkinter.filedialog
from openpyxl import Workbook


def export_docx():
    con = sqlite3.connect('polyclinic.db')
    cur = con.cursor()

    document = docx.Document()

    document.add_heading('Polyclinic Database', 0)

    document.add_heading("Doctor's table", level=1)
    doctor_table = document.add_table(rows=1, cols=6)
    doctor_cells = doctor_table.rows[0].cells
    doctor_cells[0].text = 'ID'
    doctor_cells[1].text = 'Full name'
    doctor_cells[2].text = 'Speciality'
    doctor_cells[3].text = 'Appointment cost'
    doctor_cells[4].text = 'Payroll percentage'
    doctor_cells[5].text = 'Salary'
    for row in cur.execute('SELECT * FROM doctor'):
        row_cells = doctor_table.add_row().cells
        row_cells[0].text = str(row[0])
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = str(row[3])
        row_cells[4].text = str(row[4])
        row_cells[5].text = str(row[5])
    document.add_page_break()

    document.add_heading("Patient's table", level=1)
    patient_table = document.add_table(rows=1, cols=6)
    patient_cells = patient_table.rows[0].cells
    patient_cells[0].text = 'ID'
    patient_cells[1].text = 'Full name'
    patient_cells[2].text = 'Date of birth'
    patient_cells[3].text = 'Address'
    for row in cur.execute('SELECT * FROM patient'):
        row_cells = patient_table.add_row().cells
        row_cells[0].text = str(row[0])
        row_cells[1].text = row[1]
        row_cells[2].text = str(row[2])
        row_cells[3].text = row[3]
    document.add_page_break()

    document.add_heading("Receipt's table", level=1)
    receipt_table = document.add_table(rows=1, cols=6)
    receipt_cells = receipt_table.rows[0].cells
    receipt_cells[0].text = 'ID'
    receipt_cells[1].text = 'Date'
    receipt_cells[2].text = 'Patient ID'
    receipt_cells[3].text = 'Doctor ID'
    for row in cur.execute('SELECT * FROM receipt'):
        row_cells = receipt_table.add_row().cells
        row_cells[0].text = str(row[0])
        row_cells[1].text = row[1]
        row_cells[2].text = str(row[2])
        row_cells[3].text = str(row[3])

    document.save(tkinter.filedialog.askdirectory()+'/Polyclinic_Database.docx')


def export_xlsx():
    con = sqlite3.connect('polyclinic.db')
    cur = con.cursor()
    wb = Workbook()
    ws1 = wb.create_sheet("Doctor's")
    ws2 = wb.create_sheet("Patient's")
    ws3 = wb.create_sheet("Receipt's")
    i = 1
    for row in cur.execute('SELECT * FROM doctor'):
        ws1.cell(row=i, column=1, value=row[0])
        ws1.cell(row=i, column=2, value=row[1])
        ws1.cell(row=i, column=3, value=row[2])
        ws1.cell(row=i, column=4, value=row[3])
        ws1.cell(row=i, column=5, value=row[4])
        ws1.cell(row=i, column=6, value=row[5])
        i += 1
    i = 1
    for row in cur.execute('SELECT * FROM patient'):
        ws2.cell(row=i, column=1, value=row[0])
        ws2.cell(row=i, column=2, value=row[1])
        ws2.cell(row=i, column=3, value=row[2])
        ws2.cell(row=i, column=4, value=row[3])
        i += 1
    i = 1
    for row in cur.execute('SELECT * FROM receipt'):
        ws3.cell(row=i, column=1, value=row[0])
        ws3.cell(row=i, column=2, value=row[1])
        ws3.cell(row=i, column=3, value=row[2])
        ws3.cell(row=i, column=4, value=row[3])
        i += 1

    wb.save(tkinter.filedialog.askdirectory() + '/Polyclinic_Database.xlsx')